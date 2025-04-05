import os
import json
import PyPDF2
import docx
from tqdm import tqdm
import logging
import re
from typing import List, Dict, Any

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # Precompile regex for better performance
        self.sentence_end_re = re.compile(r'(?<=[.!?])\s+')

    def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """Improved with timeout protection and better error handling"""
        try:
            logger.info(f"Processing document: {file_path}")
            
            if file_path.lower().endswith('.pdf'):
                text = self._read_pdf(file_path)
            elif file_path.lower().endswith('.docx'):
                text = self._read_docx(file_path)
            else:
                logger.warning(f"Unsupported file format: {file_path}")
                return []

            chunks = self._create_chunks_safe(text)
            return self._create_records(chunks, file_path)
            
        except Exception as e:
            logger.error(f"Critical error processing {file_path}: {str(e)}", exc_info=True)
            return []

    def _create_chunks_safe(self, text: str, timeout_sec: int = 30) -> List[str]:
        """Wrapper with timeout protection"""
        from multiprocessing import Process, Queue
        def _chunk_worker(q, text, chunk_size, chunk_overlap):
            try:
                q.put(self._create_chunks(text, chunk_size, chunk_overlap))
            except Exception as e:
                q.put(e)

        q = Queue()
        p = Process(target=_chunk_worker, 
                  args=(q, text, self.chunk_size, self.chunk_overlap))
        p.start()
        p.join(timeout=timeout_sec)
        
        if p.is_alive():
            p.terminate()
            p.join()
            logger.warning("Chunking timed out, using simple chunking")
            return self._simple_chunking(text)
            
        result = q.get()
        if isinstance(result, Exception):
            logger.warning(f"Chunking failed: {result}, using simple chunking")
            return self._simple_chunking(text)
        return result

    def _create_chunks(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """Optimized chunking with sentence boundary detection"""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # Only attempt sentence boundary detection if we're mid-text
            if end < text_length:
                # Look for the nearest sentence end within 100 chars
                search_start = max(start, end - 100)
                match = self.sentence_end_re.search(text, search_start, end + 100)
                if match:
                    end = match.end()
            
            chunk = text[start:end].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
                
            start = max(start + chunk_size - chunk_overlap, end - chunk_overlap)
            
        return chunks

    def _simple_chunking(self, text: str) -> List[str]:
        """Fallback chunking without sentence detection"""
        return [text[i:i+self.chunk_size].strip() 
               for i in range(0, len(text), self.chunk_size - self.chunk_overlap)]

    def _create_records(self, chunks: List[str], file_path: str) -> List[Dict[str, Any]]:
        """Create document records with metadata"""
        return [{
            'content': chunk,
            'metadata': {
                'source': file_path,
                'chunk_id': i,
                'file_type': os.path.splitext(file_path)[1][1:],
                'total_chunks': len(chunks)
            }
        } for i, chunk in enumerate(chunks)]

    # Rest of your methods (_read_pdf, _read_docx, process_directory) remain the same
    # ...
    def process_directory(self, directory_path: str, output_dir: str) -> None:
        """
        Process all documents in a directory.
        
        Args:
            directory_path: Path to the directory containing documents
            output_dir: Path to save processed chunks
        """
        os.makedirs(output_dir, exist_ok=True)
        
        all_chunks = []
        
        # Get all PDF and DOCX files in the directory
        document_files = [
            os.path.join(directory_path, f) for f in os.listdir(directory_path)
            if f.lower().endswith(('.pdf', '.docx'))
        ]
        
        logger.info(f"Found {len(document_files)} document files to process")
        
        for file_path in tqdm(document_files, desc="Processing documents"):
            chunks = self.process_document(file_path)
            all_chunks.extend(chunks)
            
            # Save chunks for individual file
            file_name = os.path.basename(file_path)
            output_file = os.path.join(output_dir, f"{os.path.splitext(file_name)[0]}_chunks.json")
            
            with open(output_file, 'w') as f:
                json.dump(chunks, f, indent=2)
        
        # Save all chunks to a single file
        all_chunks_file = os.path.join(output_dir, "all_document_chunks.json")
        with open(all_chunks_file, 'w') as f:
            json.dump(all_chunks, f, indent=2)
        
        logger.info(f"Processed {len(all_chunks)} total chunks from {len(document_files)} documents")
    
    def _read_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
        
        return text
    
    def _read_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
        
        return text    

















# import os
# import json
# import PyPDF2
# import docx
# from tqdm import tqdm
# import logging
# from typing import List, Dict, Any, Tuple

# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# class DocumentProcessor:
#     """
#     Processes PDF and DOCX documents to extract text content and create chunks.
#     """
    
#     def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
#         """
#         Initialize the document processor.
        
#         Args:
#             chunk_size: Size of each text chunk
#             chunk_overlap: Overlap between consecutive chunks
#         """
#         self.chunk_size = chunk_size
#         self.chunk_overlap = chunk_overlap
    
#     def process_document(self, file_path: str) -> List[Dict[str, Any]]:
#         """
#         Process a document file (PDF or DOCX) and extract text chunks.
        
#         Args:
#             file_path: Path to the document file
            
#         Returns:
#             List of dictionaries containing chunks and metadata
#         """
#         logger.info(f"Processing document: {file_path}")
        
#         if file_path.lower().endswith('.pdf'):
#             text = self._read_pdf(file_path)
#         elif file_path.lower().endswith('.docx'):
#             text = self._read_docx(file_path)
#         else:
#             logger.warning(f"Unsupported file format: {file_path}")
#             return []
        
#         # Create chunks from the extracted text
#         chunks = self._create_chunks(text)
        
#         # Create document records with metadata
#         records = []
#         for i, chunk in enumerate(chunks):
#             records.append({
#                 'content': chunk,
#                 'metadata': {
#                     'source': file_path,
#                     'chunk_id': i,
#                     'file_type': os.path.splitext(file_path)[1][1:],
#                     'total_chunks': len(chunks)
#                 }
#             })
        
#         logger.info(f"Extracted {len(records)} chunks from {file_path}")
#         return records
    
#     def process_directory(self, directory_path: str, output_dir: str) -> None:
#         """
#         Process all documents in a directory.
        
#         Args:
#             directory_path: Path to the directory containing documents
#             output_dir: Path to save processed chunks
#         """
#         os.makedirs(output_dir, exist_ok=True)
        
#         all_chunks = []
        
#         # Get all PDF and DOCX files in the directory
#         document_files = [
#             os.path.join(directory_path, f) for f in os.listdir(directory_path)
#             if f.lower().endswith(('.pdf', '.docx'))
#         ]
        
#         logger.info(f"Found {len(document_files)} document files to process")
        
#         for file_path in tqdm(document_files, desc="Processing documents"):
#             chunks = self.process_document(file_path)
#             all_chunks.extend(chunks)
            
#             # Save chunks for individual file
#             file_name = os.path.basename(file_path)
#             output_file = os.path.join(output_dir, f"{os.path.splitext(file_name)[0]}_chunks.json")
            
#             with open(output_file, 'w') as f:
#                 json.dump(chunks, f, indent=2)
        
#         # Save all chunks to a single file
#         all_chunks_file = os.path.join(output_dir, "all_document_chunks.json")
#         with open(all_chunks_file, 'w') as f:
#             json.dump(all_chunks, f, indent=2)
        
#         logger.info(f"Processed {len(all_chunks)} total chunks from {len(document_files)} documents")
    
#     def _read_pdf(self, file_path: str) -> str:
#         """Extract text from a PDF file."""
#         text = ""
#         try:
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page_num in range(len(pdf_reader.pages)):
#                     page = pdf_reader.pages[page_num]
#                     text += page.extract_text() + "\n"
#         except Exception as e:
#             logger.error(f"Error processing PDF {file_path}: {str(e)}")
        
#         return text
    
#     def _read_docx(self, file_path: str) -> str:
#         """Extract text from a DOCX file."""
#         text = ""
#         try:
#             doc = docx.Document(file_path)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         except Exception as e:
#             logger.error(f"Error processing DOCX {file_path}: {str(e)}")
        
#         return text
    
#     def _create_chunks(self, text: str) -> List[str]:
#         """Create overlapping chunks from the text."""
#         if not text.strip():
#             return []
        
#         chunks = []
#         start = 0
        
#         while start < len(text):
#             end = min(start + self.chunk_size, len(text))
            
#             # If we're not at the end of the text, try to find a sentence break
#             if end < len(text):
#                 # Look for sentence-ending punctuation followed by a space or newline
#                 for i in range(end - 1, max(start, end - 100), -1):
#                     if text[i] in '.!?' and (i + 1 >= len(text) or text[i + 1] in ' \n'):
#                         end = i + 1
#                         break
            
#             chunks.append(text[start:end].strip())
#             start = end - self.chunk_overlap
        
#         return chunks